import os
import docx
from pydub import AudioSegment
import azure.cognitiveservices.speech as speechsdk
from tqdm import tqdm
from pdfminer.high_level import extract_text
import chardet
import docx2txt

def read_docx(docx_file_path):
    doc = docx.Document(docx_file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def read_txt(file_path):
    with open(file_path, 'rb') as f:
        result = chardet.detect(f.read())
        
    with open(file_path, 'r', encoding=result['encoding']) as file:
        text = file.read()
    return text

def read_pdf(pdf_file_path):
    text = extract_text(pdf_file_path)
    if text.strip() == "":
        raise ValueError("The input PDF is image-based. Please upload a text-based PDF.")
    return text

def read_doc(doc_file_path):
    try:
        text = docx2txt.process(doc_file_path)
        return text
    except Exception as e:
        print(e)
        return None


def read_docx(docx_file_path):
    doc = docx.Document(docx_file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def text_to_speech(text, output_filename, subscription_key, region, voice_shortname, speech_recognition_language):
    speech_config = speechsdk.SpeechConfig(subscription=subscription_key, region=region, speech_recognition_language=speech_recognition_language)
    speech_config.set_speech_synthesis_output_format(speechsdk.SpeechSynthesisOutputFormat.Audio16Khz128KBitRateMonoMp3)
    speech_config.speech_synthesis_voice_name = voice_shortname

    audio_config = speechsdk.audio.AudioOutputConfig(filename=output_filename)
    synthesizer = speechsdk.SpeechSynthesizer(speech_config=speech_config, audio_config=audio_config)

    result = synthesizer.speak_text_async(text).get()
    if result.reason == speechsdk.ResultReason.SynthesizingAudioCompleted:
        print(f"Speech synthesized to [{output_filename}]")
    elif result.reason == speechsdk.ResultReason.Canceled:
        cancellation_details = result.cancellation_details
        print(f"Speech synthesis canceled: {cancellation_details.reason}")
        if cancellation_details.reason == speechsdk.CancellationReason.Error:
            print(f"Error details: {cancellation_details.error_details}")

def docx_to_mp3(file_path, subscription_key, region, language, voice):
    _, file_extension = os.path.splitext(file_path)
    if file_extension == '.docx':
        text = read_docx(file_path)
    elif file_extension == '.txt':
        text = read_txt(file_path)
    elif file_extension == '.doc':
        text = read_doc(file_path)
    elif file_extension == '.pdf':
        text = read_pdf(file_path)
    else:
        raise ValueError("Unsupported file type.")

    split_length = 2500 if language == "zh-CN" else 2000
    texts = [text[i:i + split_length] for i in range(0, len(text), split_length)]

    input_filename_without_ext, _ = os.path.splitext(os.path.basename(file_path))
    output_filenames = []
    
    # Wrap your texts list with tqdm for a progress bar
    for i, text in enumerate(tqdm(texts, desc="Converting text to speech")):
        output_filename = f"./uploads/{input_filename_without_ext}-Azure-tts-part{i}.mp3"
        text_to_speech(text, output_filename, subscription_key, region, voice, language)
        output_filenames.append(output_filename)

    # Combine all MP3 files into one
    combined = AudioSegment.empty()
    for filename in output_filenames:
        sound = AudioSegment.from_mp3(filename)
        combined += sound

    # Save the combined audio to a file
    combined_filename = f"./uploads/{input_filename_without_ext}-Azure-tts.mp3"
    combined.export(combined_filename, format="mp3")

    # Optionally, delete the individual MP3 files
    for filename in output_filenames:
        os.remove(filename)
        
    print(f"All parts combined and saved to {combined_filename}")
    return combined_filename

